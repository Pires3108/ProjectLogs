from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document

from app.errors import ApiError
from app.ingestion.normalizers import normalize_source


def test_normalizes_markdown_preserving_meaning() -> None:
    content = b"# Planejamento\n\n- Definir escopo\n- Validar entrega\n"

    result = normalize_source(".md", content)

    assert result == "Planejamento\nDefinir escopo\nValidar entrega"


def test_normalizes_docx_paragraphs_and_tables() -> None:
    document = Document()
    document.add_heading("Decisões", level=1)
    document.add_paragraph("Publicar a primeira versão.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Responsável"
    table.cell(0, 1).text = "Equipe"
    buffer = BytesIO()
    document.save(buffer)

    result = normalize_source(".docx", buffer.getvalue())

    assert "Decisões" in result
    assert "Publicar a primeira versão." in result
    assert "Responsável | Equipe" in result


@pytest.mark.parametrize("extension", [".srt", ".vtt"])
def test_normalizes_subtitles_without_timestamps(extension: str) -> None:
    content = (
        b"WEBVTT\n\n1\n00:00:01.000 --> 00:00:03.000\nOla equipe\n\n"
        b"2\n00:00:03,500 --> 00:00:05,000\nProxima pauta\n"
    )

    result = normalize_source(extension, content)

    assert result == "Ola equipe\nProxima pauta"


def test_rejects_corrupted_docx() -> None:
    with pytest.raises(ApiError) as captured:
        normalize_source(".docx", b"isto nao e um arquivo zip")

    assert captured.value.code == "CORRUPTED_DOCUMENT"


def test_rejects_docx_with_suspicious_compression_ratio() -> None:
    buffer = BytesIO()
    with ZipFile(buffer, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "types")
        archive.writestr("word/document.xml", b"0" * 2_000_000)

    with pytest.raises(ApiError) as captured:
        normalize_source(".docx", buffer.getvalue())

    assert captured.value.code == "CORRUPTED_DOCUMENT"
