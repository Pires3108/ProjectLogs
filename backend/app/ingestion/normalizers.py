import re
from abc import ABC, abstractmethod
from io import BytesIO
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from markdown_it import MarkdownIt

from app.errors import ApiError

TIMESTAMP_LINE = re.compile(
    r"^(?:\d{2}:)?\d{2}:\d{2}[,.]\d{3}\s+-->\s+(?:\d{2}:)?\d{2}:\d{2}[,.]\d{3}.*$"
)
SEQUENCE_LINE = re.compile(r"^\d+$")
VTT_METADATA = re.compile(r"^(?:WEBVTT|NOTE(?:\s.*)?|Kind:.*|Language:.*)$", re.IGNORECASE)
HTML_TAG = re.compile(r"<[^>]+>")


def decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ApiError(
        status_code=400,
        code="INVALID_TEXT_ENCODING",
        message="Não foi possível identificar a codificação do arquivo de texto.",
    )


def clean_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line).strip()


class Normalizer(ABC):
    @abstractmethod
    def normalize(self, content: bytes) -> str:
        raise NotImplementedError


class PlainTextNormalizer(Normalizer):
    def normalize(self, content: bytes) -> str:
        return clean_text(decode_text(content))


class SubtitleNormalizer(Normalizer):
    def normalize(self, content: bytes) -> str:
        text = decode_text(content)
        result: list[str] = []
        for raw_line in text.splitlines():
            line = raw_line.strip()
            if (
                not line
                or SEQUENCE_LINE.fullmatch(line)
                or TIMESTAMP_LINE.fullmatch(line)
                or VTT_METADATA.fullmatch(line)
            ):
                continue
            line = HTML_TAG.sub("", line).strip()
            if line and (not result or result[-1] != line):
                result.append(line)
        return "\n".join(result)


class MarkdownNormalizer(Normalizer):
    def __init__(self) -> None:
        self.parser = MarkdownIt("commonmark")

    def normalize(self, content: bytes) -> str:
        tokens = self.parser.parse(decode_text(content))
        parts: list[str] = []
        for token in tokens:
            if token.type in {"inline", "fence", "code_block"} and token.content.strip():
                parts.append(token.content.strip())
        return clean_text("\n".join(parts))


class DocxNormalizer(Normalizer):
    def normalize(self, content: bytes) -> str:
        try:
            with ZipFile(BytesIO(content)) as archive:
                entries = archive.infolist()
                names = {entry.filename for entry in entries}
                total_uncompressed = sum(entry.file_size for entry in entries)
                suspicious_ratio = any(
                    entry.file_size > 0 and entry.file_size / max(entry.compress_size, 1) > 1000
                    for entry in entries
                )
                if (
                    "[Content_Types].xml" not in names
                    or "word/document.xml" not in names
                    or total_uncompressed > 100 * 1024 * 1024
                    or suspicious_ratio
                ):
                    raise ValueError("unsafe or invalid docx archive")
            document = Document(BytesIO(content))
        except (BadZipFile, PackageNotFoundError, ValueError, KeyError) as exception:
            raise ApiError(
                status_code=400,
                code="CORRUPTED_DOCUMENT",
                message="O documento DOCX está corrompido ou não é válido.",
            ) from exception

        blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        return clean_text("\n".join(blocks))


NORMALIZERS: dict[str, Normalizer] = {
    ".md": MarkdownNormalizer(),
    ".docx": DocxNormalizer(),
    ".txt": PlainTextNormalizer(),
    ".vtt": SubtitleNormalizer(),
    ".srt": SubtitleNormalizer(),
}


def normalize_source(extension: str, content: bytes) -> str:
    try:
        normalizer = NORMALIZERS[extension]
    except KeyError as exception:
        raise ApiError(
            status_code=400,
            code="UNSUPPORTED_FORMAT",
            message="O formato do arquivo não possui normalizador.",
            details={"extension": extension},
        ) from exception
    return normalizer.normalize(content)
